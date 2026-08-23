#!/usr/bin/env python3
"""Word 导出与图片处理。"""
from __future__ import annotations

import gzip
import hashlib
import os
import re
import time
import zlib
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse

import requests
from docx import Document
from docx.shared import Inches, Pt

from core.utils import (
    _calc_picture_size,
    _clean_alt_for_markdown,
    _get_system_chinese_font,
    _safe_run_text,
    _set_document_default_font,
    _set_run_font,
)


class ImageRegistry:
    """图片引用注册表：用稳定占位符替代 Markdown 图片语法，避免标题/内容中的特殊字符干扰解析。"""

    _PLACEHOLDER_RE = re.compile(r'__IMG_REF__([A-Za-z0-9]+)__')

    def __init__(self):
        self._refs: dict[str, dict] = {}
        self._counter = 0

    def add(self, alt: str, url: str, local_name: str, local_path: str, relative_url: str) -> str:
        self._counter += 1
        key = f"__IMG_REF__{self._counter:06d}__"
        self._refs[key] = {
            "alt": alt,
            "url": url,
            "local_name": local_name,
            "local_path": local_path,
            "relative_url": relative_url,
            "downloaded": False,
        }
        return key

    def get(self, key: str) -> dict | None:
        return self._refs.get(key)

    def items(self):
        return self._refs.items()

    def download_all(self, max_workers: int = 5, max_retries: int = 2) -> int:
        """并发下载注册表中所有图片，更新下载状态，并返回失败图片数量。"""
        def _download_one(ref):
            key, info = ref
            success = download_image(info["url"], info["local_path"], max_retries=max_retries)
            info["downloaded"] = success
            return key, success

        if not self._refs:
            return 0

        total = len(self._refs)
        print(f"    开始并发下载 {total} 张图片（线程数 {max_workers}）...")
        completed = 0
        failed = 0
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(_download_one, ref): ref[0] for ref in self._refs.items()}
            for future in as_completed(futures):
                key, success = future.result()
                if success:
                    completed += 1
                else:
                    failed += 1
                    info = self._refs.get(key, {})
                    print(f"      [warn] 图片下载失败：{info.get('url', '')}")
        print(f"    图片下载完成：成功 {completed} 张，失败 {failed} 张")
        return failed

    def replace_failed_placeholders(self, text: str) -> str:
        """将下载失败的占位符替换为可读的失败提示。"""
        def repl(m):
            key = m.group(0)
            info = self._refs.get(key)
            if info and info.get("downloaded"):
                return key
            alt = info.get("alt", "") if info else ""
            return f"[图片加载失败: {alt}]"
        return self._PLACEHOLDER_RE.sub(repl, text)

    @classmethod
    def finditer(cls, text: str):
        return cls._PLACEHOLDER_RE.finditer(text)

    @classmethod
    def is_placeholder(cls, text: str) -> bool:
        return bool(cls._PLACEHOLDER_RE.fullmatch(text))


def _img_url_to_local_name(url: str) -> str:
    """根据图片 URL 生成本地文件名。"""
    parsed = urlparse(url)
    name = os.path.basename(parsed.path) or "img"
    if not name or "." not in name:
        name = hashlib.md5(url.encode("utf-8")).hexdigest()[:12] + ".png"
    # 保留扩展名，去掉控制字符、零宽字符及文件系统非法字符
    base, ext = os.path.splitext(name)
    base = re.sub(r'[\x00-\x1f\x7f\u200b-\u200f\ufeff]', '', base)
    base = re.sub(r'[\\/:*?"<>|\n\r\t()\[\]]', "_", base)
    base = re.sub(r'[_]+', "_", base).strip("_")
    if not ext:
        ext = ".png"
    ext = ext.lower()
    # 限制扩展名为常见图片格式，防止异常 URL 路径中的后缀被滥用
    if ext not in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".svg"):
        ext = ".png"
    return (base or "img") + ext


def _decode_response_body(data: bytes) -> bytes:
    """按魔数判断响应是否被压缩，解压出原始图片字节；无法识别则原样返回。"""
    if not data:
        return data
    # gzip 魔数
    if data[:2] == b"\x1f\x8b":
        try:
            return gzip.decompress(data)
        except Exception:
            return data
    # zlib/deflate 魔数（0x78 0x01 / 0x78 0x9c / 0x78 0xda）
    if len(data) > 2 and data[0] == 0x78 and data[1] in (0x01, 0x5E, 0x9C, 0xDA):
        try:
            return zlib.decompress(data)
        except Exception:
            return data
    return data


def download_image(url: str, local_path: str, max_retries: int = 2) -> bool:
    """下载单张图片，支持失败重试。

    使用 Request `Accept-Encoding: identity` + stream 原始字节读取，
    绕开 requests 对 CDN 返回 gzip/deflate 流时的自动解压（该解压可能因
    头不匹配抛 `zlib.error: incorrect header check` 而中断整个下载）。
    """
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36"
        ),
        # 请求未压缩响应，避免 CDN 返回压缩流时历 requests 自动解压崩溃
        "Accept-Encoding": "identity",
    }
    for attempt in range(max_retries + 1):
        try:
            resp = requests.get(url, timeout=20, headers=headers, stream=True)
            if resp.status_code != 200:
                print(f"      [warn] 下载图片失败 {url}: HTTP {resp.status_code} (尝试 {attempt + 1}/{max_retries + 1})")
            else:
                data = resp.raw.read(decode_content=False) or resp.content
                data = _decode_response_body(data)
                if data:
                    with open(local_path, "wb") as f:
                        f.write(data)
                    return True
                print(f"      [warn] 下载图片失败 {url}: 空响应体 (尝试 {attempt + 1}/{max_retries + 1})")
        except Exception as e:
            print(f"      [warn] 下载图片失败 {url}: {e} (尝试 {attempt + 1}/{max_retries + 1})")
        if attempt < max_retries:
            time.sleep(1.5 * (attempt + 1))
    return False


def download_images_in_text(text: str, images_dir: str, relative_prefix: str,
                            registry: ImageRegistry = None) -> tuple:
    """
    将 text 中所有 Markdown 图片注册到 ImageRegistry，并替换为稳定的占位符。
    图片实际下载由 ImageRegistry.download_all() 统一并发执行。
    relative_prefix 用于 Markdown 中的图片相对路径，如 "images/作业标题"。
    返回 (新文本, 图片注册表 ImageRegistry)。
    """
    if registry is None:
        registry = ImageRegistry()
    if not text:
        return text, registry

    os.makedirs(images_dir, exist_ok=True)
    # URL -> local_name 缓存，避免同一作业内重复注册同一图片
    registered: dict[str, str] = {}

    def repl(m):
        alt = _clean_alt_for_markdown(m.group(1))
        url = m.group(2)
        local_name = registered.get(url)
        if local_name is None:
            local_name = _img_url_to_local_name(url)
            registered[url] = local_name
        local_path = os.path.join(images_dir, local_name)
        relative_url = f"{relative_prefix}/{local_name}"
        return registry.add(alt, url, local_name, local_path, relative_url)

    new_text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', repl, text)
    return new_text, registry


def add_markdown_paragraph(doc: Document, text: str, images_base_dir: str,
                           prefix: str = "", prefix_bold: bool = False, prefix_color: str = None,
                           text_color: str = None, max_img_width: float = 5.5,
                           registry: ImageRegistry = None):
    """
    将含图片占位符的文本添加为 Word 段落。
    图片会根据原始尺寸按比例缩放，避免过大或过小。
    """
    from docx.enum.text import WD_LINE_SPACING

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(6)

    if prefix:
        run = paragraph.add_run(_safe_run_text(prefix))
        _set_run_font(run, bold=prefix_bold, color=prefix_color)

    last = 0
    for m in ImageRegistry.finditer(text):
        # 图片前的文字
        if m.start() > last:
            run = paragraph.add_run(_safe_run_text(text[last:m.start()]))
            _set_run_font(run, color=text_color)

        key = m.group(0)
        ref = registry.get(key) if registry else None
        if ref is None:
            run = paragraph.add_run(f"[图片引用丢失: {key}]")
            _set_run_font(run, color=text_color)
            last = m.end()
            continue

        local_path = ref["local_path"]
        alt = ref.get("alt", "")
        if os.path.exists(local_path):
            try:
                # 图片单独成段，排版更清晰
                img_para = doc.add_paragraph()
                img_para.paragraph_format.left_indent = Inches(0.3) if prefix else Inches(0)
                img_para.paragraph_format.space_before = Pt(4)
                img_para.paragraph_format.space_after = Pt(8)
                img_run = img_para.add_run()
                width, height = _calc_picture_size(local_path, max_img_width)
                if height:
                    img_run.add_picture(local_path, width=width, height=height)
                else:
                    img_run.add_picture(local_path, width=width)
            except Exception:
                run = paragraph.add_run(f"[图片加载失败: {alt}]")
                _set_run_font(run, color=text_color)
        else:
            run = paragraph.add_run(f"[图片: {alt} {ref.get('relative_url', '')}]")
            _set_run_font(run, color=text_color)
        last = m.end()

    if last < len(text):
        run = paragraph.add_run(_safe_run_text(text[last:]))
        _set_run_font(run, color=text_color)

    return paragraph


def save_word(title: str, url: str, questions: list, base: str, images_base_dir: str,
              registry: ImageRegistry = None):
    """生成排版精美的 Word 文档，包含题目、选项、答案及内嵌图片。"""
    doc = Document()
    font_name = _get_system_chinese_font()
    _set_document_default_font(doc, font_name)

    # 页面边距
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # 标题
    heading = doc.add_heading(level=1)
    run = heading.add_run(_safe_run_text(title))
    _set_run_font(run, font_name=font_name, font_size=20, bold=True, color="2B579A")
    heading.alignment = 0
    heading.paragraph_format.space_after = Pt(10)

    # 来源
    p = doc.add_paragraph()
    run = p.add_run("来源：")
    _set_run_font(run, font_size=9, color="666666")
    run = p.add_run(_safe_run_text(url))
    _set_run_font(run, font_size=9, color="666666")
    p.paragraph_format.space_after = Pt(2)

    # 总题数
    p = doc.add_paragraph()
    run = p.add_run(f"共 {len(questions)} 题")
    _set_run_font(run, font_size=10, color="888888")
    p.paragraph_format.space_after = Pt(16)

    for q_idx, q in enumerate(questions, 1):
        # 题号与题型
        qheading = doc.add_heading(level=2)
        qtype = _safe_run_text(q.get('type', '未知'))
        run = qheading.add_run(f"第 {q.get('index', q_idx)} 题  ({qtype})")
        _set_run_font(run, font_name=font_name, font_size=13, bold=True, color="2B579A")
        qheading.paragraph_format.space_before = Pt(14)
        qheading.paragraph_format.space_after = Pt(6)

        # 题干
        add_markdown_paragraph(
            doc, q.get("title", ""), images_base_dir,
            prefix="题干：", prefix_bold=True, prefix_color="2B579A",
            max_img_width=5.5, registry=registry
        )

        # 选项（不自动编号，保留原始内容）
        if q.get("options"):
            p = doc.add_paragraph()
            run = p.add_run("选项：")
            _set_run_font(run, bold=True, color="2B579A")
            for opt in q["options"]:
                add_markdown_paragraph(
                    doc, opt, images_base_dir,
                    prefix="• ", prefix_bold=True,
                    max_img_width=5.0, registry=registry
                )

        # 答案
        add_markdown_paragraph(
            doc, q.get("answer") or "未识别", images_base_dir,
            prefix="答案：", prefix_bold=True, prefix_color="C00000",
            text_color="C00000", max_img_width=5.5, registry=registry
        )

        # 题目分隔（空行）
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(10)
        sep.paragraph_format.space_after = Pt(6)

    doc_path = f"{base}.docx"
    try:
        doc.save(doc_path)
        print(f"    已保存 Word：{doc_path}")
    except Exception as e:
        print(f"    [error] 保存 Word 失败：{e}")
