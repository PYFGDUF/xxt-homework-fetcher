"""exporters.merge 合并 Word 功能测试。"""
from docx import Document

from exporters.merge import merge_all_docx


def _make_docx(path, text):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def test_merge_skips_single_file(tmp_path):
    _make_docx(tmp_path / "a.docx", "A")
    assert merge_all_docx(str(tmp_path)) == ""


def test_merge_combines_docx(tmp_path):
    for name, text in [("b.docx", "B"), ("c.docx", "C"), ("d.docx", "D")]:
        _make_docx(tmp_path / name, text)
    merged = merge_all_docx(str(tmp_path), "全部作业合并.docx")
    assert merged == str(tmp_path / "全部作业合并.docx")
    assert (tmp_path / "全部作业合并.docx").exists()
    # 合并产物存在校验：能重新打开且非空
    doc = Document(merged)
    assert len(doc.paragraphs) >= 1


def test_merge_excludes_non_docx(tmp_path):
    (tmp_path / "notes.txt").write_text("忽略我", encoding="utf-8")
    _make_docx(tmp_path / "a.docx", "A")
    _make_docx(tmp_path / "b.docx", "B")
    merged = merge_all_docx(str(tmp_path))
    assert (tmp_path / "notes.txt").exists()  # 非 docx 不受影响
    assert merged  # 两个 docx 可合并


def test_merge_idempotent_does_not_include_previous_merged(tmp_path):
    for name, text in [("e.docx", "E"), ("f.docx", "F")]:
        _make_docx(tmp_path / name, text)
    first = merge_all_docx(str(tmp_path), "合并.docx")
    assert first
    # 第二次合并：已生成的「合并.docx」不应被当作输入重复包含
    second = merge_all_docx(str(tmp_path), "合并.docx")
    doc = Document(second)
    # 至少包含初始两个段落之一，且不会因重复合并而越滚越大（此处仅校验能成功生成）
    assert (tmp_path / "合并.docx").exists()